import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Inches
import datetime
import docx
import requests
import json
from docx2pdf import convert
import pythoncom
import json

# User Inputs

# market = 'florida'  # florida is the market in the url. 
#                     # all these variables are going to be going into the url
#                     # in future coding projects, the user's inputs will affect far more than the url

# num_homes = '350'
# uipt = '1'
# region_id = '32461'

# Constants
PPSF = '$/SQUARE FEET'
ADDRESS = 'ADDRESS'
BATHS = 'BATHS'
BEDS = 'BEDS'
CITY = 'CITY'
STATE = 'STATE OR PROVINCE'
ZIP = 'ZIP OR POSTAL CODE'
PRICE = 'PRICE'
URL = 'URL'
SQFT = 'SQUARE FEET'
PRICE_90 = "ADJUSTED PRICE"
PPSF_90 = "ADJUSTED $/SQUARE FEET"
STATUS = 'STATUS'

def calculate_percentage_difference(target_number, number_set):
    """Calculate the percentage difference of a target number from the mean of a set."""
    if len(number_set) == 0:
        average = 0
    else:
        average = sum(number_set) / len(number_set)
    difference = target_number - average
    absolute_difference = abs(difference)
    percentage = (absolute_difference / average) * 100
    return percentage

def prepare_data(json_data):
    """Prepare data from the JSON source."""
    curated_data = []
    for d in json_data['payload']['homes']:
        if all(['value' in d.get(key, {}) for key in ['price', 'sqFt', 'streetLine', 'pricePerSqFt']]):
            curated_data.append({
                'STATUS': d.get('mlsStatus', 'NaN'),
                'BEDS': d.get('beds', 0),
                'BATHS': d.get('baths', 0),
                'CITY': d.get('city', 0),
                'STATE': d.get('state', 'NaN'),
                'ZIP': d.get('zip', 0),
                'PRICE': d.get('price', {}).get('value', 0),
                'PPSF': d.get('pricePerSqFt', {}).get('value', 0),
                'SQFT': d.get('sqFt', {}).get('value', 0),
                'ADDRESS': d.get('streetLine', {}).get('value', 'NaN'),
                'URL': 'https://www.redfin.com' + d.get('url', '')
            })
    return curated_data

def process_data_for_document(data):
    """Process and filter the data for the document."""
    # Save the new_data to a JSON file
    with open('processed_data.json', 'w') as json_file:
        json.dump(data, json_file, indent=4)
        
    ppsf_nums = [d['PPSF'] for d in data if d['PPSF'] is not None]
    mean_ppsf = sum(ppsf_nums) / len(ppsf_nums)

    new_data = []
    for d in data:
        percentage_below_mean = round(calculate_percentage_difference(d['PPSF'], ppsf_nums), 2)
        market_status = 'above' if d['PPSF'] > mean_ppsf else 'below'
        new_data.append({
            STATUS: d['STATUS'], 
            ADDRESS: f"{d['ADDRESS']}, {d['CITY']}, {d['STATE']} {d['ZIP']}  . . . {percentage_below_mean}% {market_status} market value.",
            PRICE: d['PRICE'], 
            PRICE_90: round(d['PRICE'] * 0.9, 2),
            SQFT: d['SQFT'], 
            PPSF: d['PPSF'], 
            PPSF_90: round(d['PPSF'] * 0.9, 2), 
            BEDS: d['BEDS'], 
            BATHS: d['BATHS'], 
            URL: d['URL']
        })
    return new_data

def generate_document(data, zipcode, total_listings, total_homes, max_ppsf, min_ppsf, max_price, min_price, mean_ppsf, mean_price):
    """Generate a DOCX document with the data."""
    doc = docx.Document()

    # Set margins
    margin_size = 0.25
    for section in doc.sections:
        section.top_margin = Inches(margin_size)
        section.bottom_margin = Inches(margin_size)
        section.left_margin = Inches(margin_size)
        section.right_margin = Inches(margin_size)

    area_label = 'area' if isinstance(zipcode, str) else 'zipcode'
    heading = doc.add_heading(f"{len(data)} Undervalued Properties in {zipcode}", level=0)
    heading.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

    # Date for the document
    formatted_date = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    doc.add_paragraph(f"There are {len(data)} underpriced properties in the {zipcode} {area_label} as of {formatted_date}."
                      "Use the links provided to go to Redfin.com to find the listing agent's contact information.")

    doc.add_paragraph("")

    # Convert the data to DataFrame for easier handling in table
    df2 = pd.DataFrame(data)

    # Add a table to the document with the DataFrame data
    table = doc.add_table(rows=len(df2)*3+1, cols=len(df2.columns)-2)
    hdr_cells = table.rows[0].cells
    for i in range(len(df2.columns)-2):
        if i == 0:
            hdr_cells[i].text = df2.columns[i]
        else: 
            hdr_cells[i].text = df2.columns[i+1]

    # Creating custom list of indexes for data
    data_rows_list = [0] + list(range(2, len(df2.columns)-1))

    # Populate the table with data
    for i in range(len(df2*3)):
        for j in range(len(data_rows_list)):
            table.rows[i*3+1].cells[j].text = str(df2.iat[i, data_rows_list[j]])
            if j == 0:
                table.rows[i*3+2].cells[0].text = str(df2.iat[i, 1])
                table.rows[i*3+3].cells[0].text = str(df2.iat[i, len(data_rows_list)+1] + '\n\n\n' + '-'*150)

    for u in range(1, len(df2)*3+1):
        if u > 1 and u % 3 == 0:
            row_cells_1 = table.rows[u].cells
            row_cells_1[0].merge(row_cells_1[-1])
            row_cells_2 = table.rows[u-1].cells
            row_cells_2[0].merge(row_cells_2[-1])

    doc.add_page_break()

    head1 = doc.add_heading("Data background", level=1)

    p5 = doc.add_paragraph(f"Even though there are {total_listings} total listings in {zipcode}, " 
                        f"there are only {total_homes} properties in this {area_label} that have" 
                        " a valid price per square foot value and are listed as either a house or home. "
                        "The data points needed are price and area, and the listings must either be denoted as a house or a home."
                        " This is specifically for residential properties. All other listings either do"
                        " not contain valid data or are not a house or home. "
                        "All other listing types were not included.\n\n"
                        
                        "The adjusted price and the adjusted price per square foot is 90% of the "
                        "listed price. These values were included because the market is currently inflated. "
                        "Listed prices do not always reflect actual values. After speaking with a few real "
                        "estate professionals, I can safely say that the values of the homes on the market are closer to 90% of what they are listed for.\n\n"
                        "When searching for properties on Redfin.com, oftentimes properties from other zipcodes in the surrounding area are also populated."
                        " This report may contain properties in zipcodes surrounding the one listed on the first page of this document.") 

    head = doc.add_heading("Why were these properties chosen over the others?", level=1)

    p1 = doc.add_paragraph(f"These {len(data)} properties were narrowed down from a total of {total_homes}. "
    f"They were listed in ascending order based on price per square foot. "
    f"The {len(data)} properties listed above were selected because they fell within the "
    f"lower quartile of the dataset. This means that out of 100% of the {total_homes} "
    f"properties accounted for, 25% percent of the properties listed were selected, hence the {len(data)} properties."
    f" \n\n 25% of {total_homes} properties = {len(data)} properties.")

    head1 = doc.add_heading("Quick stats", level=1)

    p13 = doc.add_paragraph(f"The high and low price per square foot values in the {zipcode} {area_label} are: {max_ppsf}, {min_ppsf}")
    p14 = doc.add_paragraph(f"The high and low price of valid home listings in the {zipcode} {area_label} are: {max_price}, {min_price}")
    p15 = doc.add_paragraph(f"The average price per square foot in this {area_label} is: {mean_ppsf}")
    p16 = doc.add_paragraph(f"The average price of the valid listings is: {mean_price}")

    head = doc.add_heading(f"Total listings in the {zipcode} {area_label}:  {total_listings}", level=1)
    doc.add_page_break()

    heading1 = doc.add_heading("Backstory", level=0)
    heading1.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

    heading2 = doc.add_heading("""Here are a few reasons why using price per square foot is one of the most important metrics when purchasing or selling a property.""", level=1)
    #heading2.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")

    # here we add a paragraph explaining the advantages of using price per square foot when determining a properties value
    p2 = doc.add_paragraph("""One of the advantages of using price per square foot as a metric is its consistency. 
    Unlike other factors that can be subjective, such as the condition of a property or its location, 
    the price per square foot is a standardized measure that allows for apples-to-apples comparisons. 
    It provides a level playing field for evaluating properties of different sizes, styles, and locations, 
    and eliminates the bias that can be associated with subjective factors.

    Analyzing the price per square foot can reveal valuable insights about the market trends and dynamics. 
    By keeping track of the historical price per square foot for a particular area, investors can identify 
    patterns and trends, such as whether prices are rising or falling, or how a specific property compares 
    to the overall market. This information can help investors make informed decisions and capitalize on 
    opportunities in the market.

    Another benefit of using price per square foot is its versatility. It can be used for different types 
    of properties, including single-family homes, multi-family properties, condos, and commercial properties. 
    Investors can apply this metric to various markets, whether it's a booming urban area or a quiet 
    suburban neighborhood. This flexibility makes it a valuable tool for investors with different 
    investment strategies and goals.

    Using price per square foot to assess a property's value is not only limited to buying, but also 
    extends to selling and renting properties. For sellers, knowing the price per square foot of 
    comparable properties can help determine the optimal listing price and attract potential buyers. 
    For landlords, understanding the price per square foot of rental properties can assist in setting 
    appropriate rent levels and maximizing rental income.""")

    #p2.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")

    # choosing to not list the complete data
    doc.add_page_break()
    heading1 = doc.add_heading("Complete Data Set", level=0)
    heading1.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    heading2 = doc.add_heading(f"Here is the complete data set listing all {total_homes} properties with valid data in ascending order based on price per square foot.", level=1)
    doc.add_paragraph("")
    p2 = doc.add_paragraph("""The complete data set has not been included in order to speed up the process of creating reports and to reduce distractions to users.""")
    
    return doc

def convert_to_pdf(input_path, output_path):
    """
    Convert a Word (.docx) file to PDF using the docx2pdf library.
    """
    try:
        pythoncom.CoInitialize()
        convert(input_path, output_path)
        pythoncom.CoUninitialize()
    except Exception as e:
        print(f"Error during conversion: {e}")


# Main function to be called
def main(num_homes, uipt, region_id):
    # response = requests.get(f'https://www.redfin.com/stingray/api/gis?al=1&has_deal=false&has_dishwasher=false&has_laundry_facility=false&has_laundry_hookups=false&has_parking=false&has_pool=false&has_short_term_lease=false&include_pending_homes=false&isRentals=false&is_furnished=false&is_income_restricted=false&is_senior_living=false&market={market}&num_homes={num_homes}&ord=redfin-recommended-asc&page_number=1&region_id={region_id}&region_type=2&sf=1,3,7&status=9&travel_with_traffic=false&travel_within_region=false&uipt={uipt}&utilities_included=false&v=8')
    response = requests.get(f'https://www.redfin.com/stingray/api/gis?al=1&has_deal=false&has_dishwasher=false&has_laundry_facility=false&has_laundry_hookups=false&has_parking=false&has_pool=false&has_short_term_lease=false&include_pending_homes=false&isRentals=false&is_furnished=false&is_income_restricted=false&is_senior_living=false&num_homes={num_homes}&ord=redfin-recommended-asc&page_number=1&region_id={region_id}&region_type=2&sf=1,3,7&status=9&travel_with_traffic=false&travel_within_region=false&uipt={uipt}&utilities_included=false&v=8')

    text_ = response.text
    new_text = text_.replace('{}&&', '')
    json_data = json.loads(new_text)

    # Prepare and process data
    data = prepare_data(json_data)
    processed_data = process_data_for_document(data)
    print(processed_data)
    
    # Filtered data operations
    filtered_data = [d for d in data if all(d[key] is not None for key in ['PRICE', 'SQFT', 'PPSF', 'ADDRESS', 'CITY', 'ZIP'])]
    filtered_data.sort(key=lambda x: x['PPSF'])
    zipcode = str(filtered_data[0]['CITY']) if filtered_data else 'Unknown'
    
    # Basic calculations
    total_listings = len(data)
    total_homes = len(filtered_data)
    max_ppsf = max(d['PPSF'] for d in filtered_data)
    min_ppsf = min(d['PPSF'] for d in filtered_data)
    max_price = max(d['PRICE'] for d in filtered_data)
    min_price = min(d['PRICE'] for d in filtered_data)
    mean_ppsf = round(sum(d['PPSF'] for d in filtered_data) / len(filtered_data), 2)
    mean_price = round(sum(d['PRICE'] for d in filtered_data) / len(filtered_data), 2)

    # Extract URLs
    urls = [d[URL] for d in filtered_data]

    # Refining data to the 25th percentile
    q25 = np.percentile([d[PPSF] for d in processed_data], 25)
    data_25 = [d for d in processed_data if d[PPSF] <= q25]

    # Generate the document
    doc = generate_document(data_25, zipcode, total_listings, total_homes, max_ppsf, min_ppsf, max_price, min_price, mean_ppsf, mean_price)
    
    # Save the document as DOCX
    if type(zipcode) is str:
        doc_path = f'Undervalued_Properties.docx'
        pdf_path = f'Undervalued_Properties.pdf'
    else:
        doc_path = f'ALL_The_Undervalued_Properties.docx'
        pdf_path = f'ALL_The_Undervalued_Properties.pdf'
    doc.save(doc_path)
    
    # Convert the DOCX to PDF
    convert_to_pdf(doc_path, pdf_path)

    return zipcode, data_25


# Call the main function
if __name__ == "__main__":
    main()
